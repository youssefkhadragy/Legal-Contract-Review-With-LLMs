from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
import shutil
import os
import string
import random
import re
from docx import Document
import fitz
import io 
import uvicorn
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import json
from fastapi.responses import StreamingResponse
from io import BytesIO


# Initialize the FastAPI app
app = FastAPI()

# Global variable to store the name of the directory where the uploaded files will be saved
UPLOAD_DIR = "uploads"

# Global variable to store file path
file_path = ""  

# Initialize a Gemini LLM model using the Gemini free tier API (The model is configered to return a JSON file using the generation_config parameter)
model = genai.GenerativeModel('gemini-1.5-flash', generation_config={"response_mime_type": "application/json"})

# Regular expressions to search, using logic, for key clauses/concepts in the contract
CLAUSE_PATTERNS = {
    "Termination": r"(?i)\btermination\b",
    "Confidentiality": r"(?i)\bconfidentiality\b",
    "Payment Terms": r"(?i)\bpayment\s+terms\b",
    "Governing Law": r"(?i)\bgoverning\s+law\b",
    "Dispute Resolution": r"(?i)\bdispute\s+resolution\b",
    "Intellectual Property": r"(?i)\bintellectual\s+property\b",
    "Liability": r"(?i)\bliability\b",
    "Indemnity": r"(?i)\bindemnity\b",
    "Force Majeure": r"(?i)\bforce\s+majeure\b",
    "Amendments": r"(?i)\bamendments\b",
    "Entire Agreement": r"(?i)\bentire\s+agreement\b",
    "Severability": r"(?i)\bseverability\b",
    "Notices": r"(?i)\bnotices\b",
    "Representation and Warranties": r"(?i)\brepresentation\s+and\s+warranties\b",
    "Assignment": r"(?i)\bassigment\b",
    "Waiver": r"(?i)\bwaiver\b",
    "Subcontracting": r"(?i)\bsubcontracting\b",
    "Responsibilities": r"(?i)\bresponsibilities\b",
    "Governing Language": r"(?i)\bgoverning\s+language\b",
    "Disclosures": r"(?i)\bdisclosures\b",
    "Audit Rights": r"(?i)\baudits\s+rights\b",
    "Counterparts": r"(?i)\bcounterparts\b",
    "Headings": r"(?i)\bheadings\b",
    "Contract Duration": r"(?i)\bcontract\s+duration\b",
    "Miscellaneous": r"(?i)\bmiscellaneous\b",
    "Insurance": r"(?i)\binsurance\b",
    "Indemnification": r"(?i)\bindemnification\b",
}





# Google API key for Generative AI
GOOGLE_API_KEY = "GOOGLE_API_KEY"

# Configure the Generative AI with the API key
genai.configure(api_key=GOOGLE_API_KEY)

# The following code segment creates the upload directory if it does not exist
try:
    os.makedirs(UPLOAD_DIR, exist_ok=True)
except Exception as e:
    print(f"Error occurred while creating upload directory: {e}")

def generate_random_string(length=12):
    """
    Generates a random string of the specified length.

    Parameters:
        length (int): The length of the random string. Default is 12.

    Returns:
        str: A random string of the specified length.
    """
    letters = string.ascii_letters
    return ''.join(random.choice(letters) for i in range(length))


def read_docx_from_upload(file):
    """
    Reads the content of a docx file from bytes.

    Parameters:
        file (bytes): The content of the docx file in bytes.

    Returns:
        str: The text content of the docx file.
    """
    
    doc = Document(io.BytesIO(file))
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def read_pdf_from_upload(file):
    """
    Reads the content of a PDF file from bytes.
    
    Parameters:
        file (bytes): The content of the PDF file in bytes.
    
    Returns:
        str: The text content of the PDF file.
    """
    doc = fitz.open("pdf", io.BytesIO(file))
    text = ""
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text += page.get_text()
    
    return text

def clean_text(text):
    """
    Cleans the extracted text by removing non-printable characters, extra whitespace, and standardizing formatting. 
    (This function is applied on text extracted from PDF files only and not docx files).
    
    Parameters:
        text (str): The text to be cleaned.
    
    Returns:
        str: The cleaned text.
    """
    # Remove non-printable characters
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    
    # Replace problematic characters with standard characters (In this case smart quotes to straight quote)
    text = text.replace('\u2019', "'")  
    text = text.replace('\u201c', '"')  
    text = text.replace('\u201d', '"')  
    
    # Remove unnecessary line breaks
    text = re.sub(r'\n+', '\n', text)  
    
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Standardize headings (Assuming headings are in all caps and followed by a period)
    text = re.sub(r'(\b[A-Z ]+)\.', r'**\1**', text)
    
    # Remove extra spaces before punctuation
    text = re.sub(r'\s([?.!,])', r'\1', text)
    
    return text

def identify_clauses(text):
    """
    Identifies key clauses/concepts in the given text using the pre-defined patterns.
    (Pre-defined patterns are stored in the global variable CLAUSE_PATTERNS).

    Parameters:
        text (str): The text of the contract to search for clauses/concepts in.

    Returns:
        dict: A dictionary of identified clauses with their corresponding text.
    """
    # A dictionary to store the retrieved patterns from the text in the follwing format: { key_clause/concept : relevant_text }
    identified_clauses = {}
    
    # A loop that iterates over pairs in the CLAUSE_PATTERNS dictionary
    for clause_name, pattern in CLAUSE_PATTERNS.items():
        
        # Match stores the first text that matches a statement for every clause pattern
        match = re.search(pattern, text, re.IGNORECASE)
        
        # The following code segment extracts 4 sentences starting from the word that matched the pattern
        if match:
            clause_text = text[match.start():]
            
            stop = 0
            
            for i, char in enumerate(clause_text):
                if char == '.':
                    stop += 1
                if stop == 4:
                    end_index = i+1
                    break
            
            if end_index != -1:
                clause_text = clause_text[:end_index]
            
            # Clean the extracted clause text
            cleaned_text = clause_text.strip().replace('\n', '')
            identified_clauses[clause_name] = cleaned_text
    return identified_clauses


def apply_revisions_to_docx(docx_content: bytes, revisions: dict) -> bytes:
    """
    Applies revisions to a docx file based on the provided revisions.
    
    Parameters:
        docx_content (bytes): The binary content of the docx file.
        revisions (dict): A dictionary where keys are original text to be replaced and values are the revised text.

    Returns:
        bytes: The binary content of the updated docx file with the revisions applied.
    """
    # Load the docx file, into a document object, from its binary content
    doc = Document(BytesIO(docx_content))
    
    # Iterate through paragraphs and runs
    for paragraph in doc.paragraphs:
        for original_text, revised_text in revisions.items():
            if original_text in paragraph.text:
                paragraph.text = paragraph.text.replace(original_text, revised_text)
                continue
            
        for run in paragraph.runs:
            
            # Apply revisions by replacing original text with revised text
            for original_text, revised_text in revisions.items():
                if original_text in run.text:
                    run.text = run.text.replace(original_text, revised_text)
    
    # Save the revised document to a BytesIO stream and return the bytes of the modified document
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)  
    return output_stream.getvalue()  


@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    """
    Uploads a file and returns a mock URL and contract type (Uploaded file is stored locally in the 'upload' directory created earlier).

    Parameters:
        file (UploadFile): The file to upload.

    Returns:
        JSONResponse: The mock URL and contract type.
    """
    global file_path
    unique_filename = f"{generate_random_string()}.{file.filename.split('.')[-1]}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    mock_url = f"https://monstersgraphics.com/{unique_filename}"
    
    contract_type = "General Contract"
    if "nda" in file.filename.lower():
        contract_type = "NDA"
    elif "employment" in file.filename.lower():
        contract_type = "Employment Contract"
    
    return JSONResponse(content={"Mock_URL": mock_url, "Contract_Type": contract_type})

@app.post("/identify-clauses-logical/")
async def identify_clauses_logical(file: UploadFile = File(...)):
    """
    Identifies clauses in the uploaded file using logical pattern matching.

    Parameters:
        file (UploadFile): The file to process.

    Returns:
        JSONResponse: The identified clauses.
    """
    content = await file.read()
    text = ""
    if file.filename.endswith(".docx"):
        text = read_docx_from_upload(content)
    elif file.filename.endswith(".pdf"):
        text = clean_text(read_pdf_from_upload(content))
    else:
        return JSONResponse(content={"Error": "Unsupported file type!"}, status_code=400)
    
    clauses = identify_clauses(text)
    
    return JSONResponse(content=clauses)



@app.post("/identify-clauses-llm/")
async def identify_clauses_llm(file: UploadFile = File(...)):
    """
    Identifies clauses in the uploaded file using Gemini LLM accessed throguh Google API.

    Parameters:
        file (UploadFile): The file to process.

    Returns:
        JSONResponse: The identified clauses.
    """
    
    
    content = await file.read()
    text = ""
    
    # Call the function specific to the type of the file to extract text
    if file.filename.endswith(".docx"):
        text = read_docx_from_upload(content)
    elif file.filename.endswith(".pdf"):
        text = clean_text(read_pdf_from_upload(content))
    else:
        return JSONResponse(content={"Error": "Unsupported file type!"}, status_code=400)
    
    # Prompt
    response = model.generate_content(
        f"Find the key clauses/concepts in the contract that is given and identify the contiguous text block(s) corresponding to each key clause/concept as an expert contract manager. "
        f"Key clauses/concepts are for example but not limited to: Termination Terms, Confidentiality Agreements, Payment Terms, etc. "
        f"Here is the contract text: [Start of Contract Text] {text} [End of Contract Text]. "
        f"Answer should be less than 800 words. Generate only the output and ensure the format is as follows: {{'key_clause/concept': 'corresponding_contiguous_text_block(s)'}}"
    )    
    
    
    try:
        return JSONResponse(content=json.loads(response.text))
    except Exception as e:
        return JSONResponse(content={'Error': f'Internal server error: {e}'})

@app.post("/revise-contract-llm/")
async def revisions_suggestion(file: UploadFile = File(...)):
    """
    Suggests revisions to the uploaded contract using Gemini LLM accessed throguh Google API.

    Parameters:
        file (UploadFile): The file to process.

    Returns:
        JSONResponse: The suggested revisions.
    """
    
    content = await file.read()
    text = ""
    
    # Call the function specific to the type of the file to extract text
    if file.filename.endswith(".docx"):
        text = read_docx_from_upload(content)
    elif file.filename.endswith(".pdf"):
        text = clean_text(read_pdf_from_upload(content))
    else:
        return JSONResponse(content={"Error": "Unsupported file type!"}, status_code=400)
    
    # Prompt (The modifications are limited to two text blocks as without this limit the response was so long that it gets blocked)
    response = model.generate_content(
        f"Review the following text and suggest linguistic improvements that enhances the negotiation terms, clarity and readability. "
        f"The suggestions should be tailored to the type of text provided. "
        f"Here is the text: {text}"
        f"Generate only the output and ensure that it is in the following format: "
        f"{{'original_text_before_revision': 'text_after_revision'}}. "
        f"Ensure that each key in the output JSON is composed of the original text block, "
        f"and the value is the modified text block after revision."
        f"Modify only 2 text blocks from the whole text. The revised block should not exceed 120 words."
        ,
        safety_settings={
        HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
        HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE
    })
    
    # Perform additional text cleaning
    text = response.text.replace('\n', '').replace('\n\n', '').replace('\n', '')
    cleaned_text = re.sub(r'\n+', '', text)
    
    try:
        return JSONResponse(content=(json.loads(cleaned_text)))
    except Exception as e:
        return JSONResponse(content={'Error': f'Internal server error: {e}'})
    
@app.post("/revise-contract-llm-return-doc/")
async def revisions_suggestion_returns_doc(file: UploadFile = File(...)):
    """
    Suggests and apply revisions to the uploaded contract using Gemini LLM accessed throguh Google API.
    
    Parameters:
        file (UploadFile): The file to process.
    
    Returns:
        StreamingResponse: A new .docx revised contract file.
    """
    content = await file.read()
    text = ""
    
    # Call the function specific to the type of the file to extract text
    if file.filename.endswith(".docx"):
        text = read_docx_from_upload(content)
    elif file.filename.endswith(".pdf"):
        text = clean_text(read_pdf_from_upload(content))
    else:
        return JSONResponse(content={"Error": "Unsupported file type!"}, status_code=400)
    
    # Prompt (The modifications are limited to two text blocks as without this limit the response was so long that it gets blocked)
    response = model.generate_content(
        f"Review the following text and suggest linguistic improvements that enhance the negotiation terms, clarity, and readability. "
        f"The suggestions should be tailored to the type of text provided. "
        f"Here is the text: {text}"
        f"Generate only the output and ensure that it is in the following format: "
        f"{{'original_text_before_revision': 'text_after_revision'}}. "
        f"Ensure that each key in the output JSON is composed of the original text block, "
        f"and the value is the modified text block after revision."
        f"Modify only 2 text blocks from the whole text. The revised block should not exceed 120 words."
        ,
        
        # Disable some thresholds that prevented Gemini from revising the contract 
        safety_settings={
        HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
        HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE
    })
    
    
    # Perform additional text cleaning
    text = response.text.replace('\n', '').replace('\n\n', '').replace('\n', '')
    cleaned_text = re.sub(r'\n+', '', text)
    
    try:
        revisions = json.loads(cleaned_text)
        
        # Apply revisions and create a new docx
        new_doc_content = apply_revisions_to_docx(content, revisions)
        
        # Return the new .docx file
        return StreamingResponse(
            BytesIO(new_doc_content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=new_contract.docx"}
        )
    
    except Exception as e:
        return JSONResponse(content={'Error': f'Internal server error: {e}'})

if __name__ == "__main__":
    # Run the FastAPI app with Uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)